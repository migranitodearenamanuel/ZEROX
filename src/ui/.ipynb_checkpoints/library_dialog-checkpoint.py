#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BIBLIOTECA DE CONOCIMIENTO DE ZEROX
===================================
Esta ventana permite cargar libros de trading para que la IA aprenda.
Mientras más libros le des, más inteligente se vuelve.
Soporta: PDF, EPUB, AZW3, MOBI, TXT, DOCX
"""

import os      # Para trabajar con archivos y carpetas
import json    # Para guardar información de los libros
import shutil  # Para copiar archivos
from PyQt6.QtWidgets import (
    QDialog,         # Ventana de diálogo
    QVBoxLayout,     # Organiza elementos verticalmente
    QHBoxLayout,     # Organiza elementos horizontalmente
    QPushButton,     # Botones
    QLabel,          # Etiquetas de texto
    QListWidget,     # Lista de elementos
    QListWidgetItem, # Elementos de la lista
    QFileDialog,     # Diálogo para seleccionar archivos
    QProgressBar,    # Barra de progreso
    QTextEdit,       # Área de texto
    QGroupBox,       # Grupo de elementos
    QMessageBox,     # Mensajes emergentes
    QSplitter,       # Divide la ventana en secciones
    QWidget          # Widget base
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer  # Funcionalidades del núcleo
from PyQt6.QtGui import QFont, QIcon, QDragEnterEvent, QDropEvent  # Elementos gráficos

# Importar las librerías para leer diferentes tipos de libros
try:
    import PyPDF2              # Para leer PDFs
    import ebooklib           # Para leer EPUBs
    from ebooklib import epub # Específico para EPUBs
    from docx import Document # Para leer archivos Word
    import mobi              # Para leer archivos MOBI (Kindle)
except ImportError as e:
    print(f"⚠️ Falta instalar librerías para leer libros: {e}")


class BookProcessor(QThread):
    """
    PROCESADOR DE LIBROS EN SEGUNDO PLANO
    Procesa los libros sin congelar la interfaz.
    Es como un trabajador que lee el libro mientras tú sigues usando el programa.
    """
    
    # Señales para comunicarse con la interfaz principal
    progress = pyqtSignal(int)      # Envía el progreso (0-100%)
    status = pyqtSignal(str)        # Envía mensajes de estado
    finished = pyqtSignal(dict)     # Envía los datos cuando termina
    error = pyqtSignal(str)         # Envía errores si algo falla
    
    def __init__(self, file_path):
        """
        Constructor del procesador
        
        Args:
            file_path: Ruta del archivo del libro a procesar
        """
        super().__init__()
        self.file_path = file_path
        
    def run(self):
        """
        Función principal que procesa el libro
        Se ejecuta en un hilo separado para no bloquear la interfaz
        """
        try:
            # Informar que estamos procesando
            self.status.emit(f"📖 Procesando: {os.path.basename(self.file_path)}")
            
            # Detectar el tipo de archivo por su extensión
            ext = os.path.splitext(self.file_path)[1].lower()
            content = ""  # Aquí guardaremos el texto del libro
            
            # Información básica del libro
            metadata = {
                "title": os.path.basename(self.file_path),  # Nombre del archivo
                "type": ext,                                # Tipo de archivo
                "size": os.path.getsize(self.file_path),
                "path": self.file_path
            }
            
            # Procesar según el tipo
            if ext == '.pdf':
                content = self._process_pdf()
            elif ext == '.epub':
                content = self._process_epub()
            elif ext in ['.azw3', '.mobi']:
                content = self._process_mobi()
            elif ext == '.txt':
                content = self._process_txt()
            elif ext == '.docx':
                content = self._process_docx()
            else:
                self.error.emit(f"Formato no soportado: {ext}")
                return
                
            # Extraer información importante
            metadata["content_length"] = len(content)
            metadata["summary"] = self._extract_summary(content)
            metadata["key_concepts"] = self._extract_key_concepts(content)
            
            self.progress.emit(100)
            self.status.emit("¡Libro procesado con éxito!")
            self.finished.emit(metadata)
            
        except Exception as e:
            self.error.emit(f"Error procesando libro: {str(e)}")
            
    def _process_pdf(self):
        """Procesa archivos PDF"""
        content = ""
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for i, page in enumerate(pdf_reader.pages):
                    content += page.extract_text() + "\n"
                    progress = int((i + 1) / total_pages * 80)
                    self.progress.emit(progress)
                    
        except Exception as e:
            self.error.emit(f"Error leyendo PDF: {str(e)}")
            
        return content
        
    def _process_epub(self):
        """Procesa archivos EPUB"""
        content = ""
        try:
            book = epub.read_epub(self.file_path)
            items = list(book.get_items())
            total_items = len(items)
            
            for i, item in enumerate(items):
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    content += item.get_content().decode('utf-8', errors='ignore') + "\n"
                    
                progress = int((i + 1) / total_items * 80)
                self.progress.emit(progress)
                
        except Exception as e:
            self.error.emit(f"Error leyendo EPUB: {str(e)}")
            
        return content
        
    def _process_mobi(self):
        """Procesa archivos MOBI/AZW3"""
        content = ""
        try:
            # Usar la librería mobi para extraer contenido
            tempdir, filepath = mobi.extract(self.file_path)
            
            # Leer el HTML extraído
            for root, dirs, files in os.walk(tempdir):
                for file in files:
                    if file.endswith('.html'):
                        with open(os.path.join(root, file), 'r', encoding='utf-8') as f:
                            content += f.read() + "\n"
                            
            # Limpiar archivos temporales
            shutil.rmtree(tempdir)
            
        except Exception as e:
            self.error.emit(f"Error leyendo MOBI/AZW3: {str(e)}")
            
        return content
        
    def _process_txt(self):
        """Procesa archivos TXT"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            self.error.emit(f"Error leyendo TXT: {str(e)}")
            return ""
            
    def _process_docx(self):
        """Procesa archivos DOCX"""
        content = ""
        try:
            doc = Document(self.file_path)
            total_paragraphs = len(doc.paragraphs)
            
            for i, paragraph in enumerate(doc.paragraphs):
                content += paragraph.text + "\n"
                progress = int((i + 1) / total_paragraphs * 80)
                self.progress.emit(progress)
                
        except Exception as e:
            self.error.emit(f"Error leyendo DOCX: {str(e)}")
            
        return content
        
    def _extract_summary(self, content):
        """Extrae un resumen del contenido"""
        # Tomar las primeras 500 palabras
        words = content.split()[:500]
        return ' '.join(words) + "..."
        
    def _extract_key_concepts(self, content):
        """Extrae conceptos clave del libro"""
        # Lista de palabras clave de trading
        keywords = [
            'trading', 'inversión', 'análisis técnico', 'análisis fundamental',
            'gestión de riesgo', 'stop loss', 'take profit', 'tendencia',
            'soporte', 'resistencia', 'volumen', 'volatilidad', 'estrategia',
            'indicadores', 'RSI', 'MACD', 'media móvil', 'fibonacci',
            'velas japonesas', 'patrones', 'breakout', 'pullback'
        ]
        
        found_concepts = []
        content_lower = content.lower()
        
        for keyword in keywords:
            if keyword.lower() in content_lower:
                found_concepts.append(keyword)
                
        return found_concepts[:10]  # Máximo 10 conceptos


class LibraryDialog(QDialog):
    """Diálogo para gestionar la biblioteca de libros"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.library_path = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'library')
        self.processed_books = []
        self._init_ui()
        self._load_library()
        
    def _init_ui(self):
        """Inicializa la interfaz"""
        self.setWindowTitle("📚 Biblioteca de Conocimiento - ZEROX")
        self.setFixedSize(900, 600)
        self.setStyleSheet("""
            QDialog {
                background-color: #0a0e1a;
                color: #ffffff;
            }
            QLabel {
                color: #8899aa;
                font-size: 12px;
            }
            QListWidget {
                background-color: #1a1f2e;
                border: 1px solid #2a3f5f;
                color: #ffffff;
                font-size: 12px;
            }
            QListWidget::item:selected {
                background-color: #2a3f5f;
            }
            QPushButton {
                background-color: #2a3f5f;
                color: #ffffff;
                border: none;
                padding: 10px 20px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #3a4f6f;
            }
            QTextEdit {
                background-color: #1a1f2e;
                border: 1px solid #2a3f5f;
                color: #ffffff;
                font-size: 12px;
            }
            QProgressBar {
                background-color: #1a1f2e;
                border: 1px solid #2a3f5f;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #00d4ff;
            }
            QGroupBox {
                border: 1px solid #2a3f5f;
                margin-top: 10px;
                padding-top: 10px;
                color: #00d4ff;
                font-weight: bold;
            }
        """)
        
        # Layout principal
        layout = QVBoxLayout()
        
        # Título
        title = QLabel("📚 BIBLIOTECA DE CONOCIMIENTO")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size: 20px; color: #00d4ff; font-weight: bold; padding: 10px;")
        layout.addWidget(title)
        
        # Descripción
        desc = QLabel(
            "Arrastra libros aquí para que ZEROX aprenda estrategias avanzadas.\n"
            "Formatos soportados: PDF, EPUB, AZW3, MOBI, TXT, DOCX"
        )
        desc.setAlignment(Qt.AlignmentFlag.AlignCenter)
        desc.setStyleSheet("color: #8899aa; padding: 5px;")
        layout.addWidget(desc)
        
        # Splitter para dividir la vista
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Panel izquierdo - Lista de libros
        left_widget = QWidget()
        left_layout = QVBoxLayout()
        
        # Lista de libros
        self.book_list = QListWidget()
        self.book_list.setAcceptDrops(True)
        self.book_list.dragEnterEvent = self._drag_enter_event
        self.book_list.dropEvent = self._drop_event
        self.book_list.itemClicked.connect(self._on_book_selected)
        left_layout.addWidget(QLabel("📖 Libros Cargados:"))
        left_layout.addWidget(self.book_list)
        
        # Botones
        btn_layout = QHBoxLayout()
        
        self.btn_add = QPushButton("➕ Añadir Libros")
        self.btn_add.clicked.connect(self._add_books)
        btn_layout.addWidget(self.btn_add)
        
        self.btn_remove = QPushButton("🗑️ Eliminar")
        self.btn_remove.clicked.connect(self._remove_book)
        self.btn_remove.setEnabled(False)
        btn_layout.addWidget(self.btn_remove)
        
        left_layout.addLayout(btn_layout)
        
        # Estadísticas
        self.stats_label = QLabel("📊 Total: 0 libros | 0 MB")
        self.stats_label.setStyleSheet("color: #667788; font-size: 11px; padding: 5px;")
        left_layout.addWidget(self.stats_label)
        
        left_widget.setLayout(left_layout)
        splitter.addWidget(left_widget)
        
        # Panel derecho - Detalles del libro
        right_widget = QWidget()
        right_layout = QVBoxLayout()
        
        # Información del libro
        info_group = QGroupBox("📋 Información del Libro")
        info_layout = QVBoxLayout()
        
        self.book_info = QTextEdit()
        self.book_info.setReadOnly(True)
        self.book_info.setMaximumHeight(200)
        info_layout.addWidget(self.book_info)
        
        info_group.setLayout(info_layout)
        right_layout.addWidget(info_group)
        
        # Conceptos extraídos
        concepts_group = QGroupBox("🎯 Conceptos Clave Detectados")
        concepts_layout = QVBoxLayout()
        
        self.concepts_list = QListWidget()
        self.concepts_list.setMaximumHeight(150)
        concepts_layout.addWidget(self.concepts_list)
        
        concepts_group.setLayout(concepts_layout)
        right_layout.addWidget(concepts_group)
        
        # Progreso
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        right_layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel("")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setStyleSheet("color: #00d4ff;")
        right_layout.addWidget(self.status_label)
        
        right_layout.addStretch()
        right_widget.setLayout(right_layout)
        splitter.addWidget(right_widget)
        
        # Establecer proporciones del splitter
        splitter.setSizes([400, 500])
        
        layout.addWidget(splitter)
        
        # Botón cerrar
        self.btn_close = QPushButton("✅ Cerrar")
        self.btn_close.clicked.connect(self.accept)
        self.btn_close.setStyleSheet("""
            QPushButton {
                background-color: #00d4ff;
                color: #0a0e1a;
                font-size: 16px;
                padding: 12px;
            }
            QPushButton:hover {
                background-color: #00a0cc;
            }
        """)
        layout.addWidget(self.btn_close)
        
        self.setLayout(layout)
        
    def _drag_enter_event(self, event: QDragEnterEvent):
        """Maneja el evento de arrastrar archivos"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            
    def _drop_event(self, event: QDropEvent):
        """Maneja el evento de soltar archivos"""
        files = []
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if os.path.isfile(file_path):
                ext = os.path.splitext(file_path)[1].lower()
                if ext in ['.pdf', '.epub', '.azw3', '.mobi', '.txt', '.docx']:
                    files.append(file_path)
                    
        if files:
            self._process_files(files)
            
    def _add_books(self):
        """Abre diálogo para seleccionar libros"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Seleccionar Libros",
            "",
            "Libros (*.pdf *.epub *.azw3 *.mobi *.txt *.docx);;Todos (*.*)"
        )
        
        if files:
            self._process_files(files)
            
    def _process_files(self, files):
        """Procesa los archivos seleccionados"""
        for file_path in files:
            # Verificar si ya está procesado
            if any(book['path'] == file_path for book in self.processed_books):
                QMessageBox.information(
                    self,
                    "Libro Duplicado",
                    f"{os.path.basename(file_path)} ya está en la biblioteca."
                )
                continue
                
            # Copiar archivo a la biblioteca
            ext = os.path.splitext(file_path)[1].lower()
            dest_folder = os.path.join(self.library_path, ext[1:])  # Quitar el punto
            os.makedirs(dest_folder, exist_ok=True)
            
            dest_path = os.path.join(dest_folder, os.path.basename(file_path))
            shutil.copy2(file_path, dest_path)
            
            # Procesar el libro
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            
            processor = BookProcessor(dest_path)
            processor.progress.connect(self.progress_bar.setValue)
            processor.status.connect(self.status_label.setText)
            processor.finished.connect(self._on_book_processed)
            processor.error.connect(self._on_processing_error)
            processor.start()
            
    def _on_book_processed(self, metadata):
        """Callback cuando un libro se procesa exitosamente"""
        self.processed_books.append(metadata)
        
        # Añadir a la lista
        item = QListWidgetItem(f"📖 {metadata['title']}")
        item.setData(Qt.ItemDataRole.UserRole, metadata)
        self.book_list.addItem(item)
        
        # Actualizar estadísticas
        self._update_stats()
        
        # Guardar biblioteca
        self._save_library()
        
        # Ocultar progreso
        QTimer.singleShot(2000, lambda: self.progress_bar.setVisible(False))
        
        # Notificar a la IA
        self._notify_ai(metadata)
        
    def _on_processing_error(self, error):
        """Maneja errores de procesamiento"""
        QMessageBox.critical(self, "Error", error)
        self.progress_bar.setVisible(False)
        self.status_label.setText("")
        
    def _on_book_selected(self, item):
        """Muestra información del libro seleccionado"""
        metadata = item.data(Qt.ItemDataRole.UserRole)
        
        # Mostrar información
        info_text = f"""
        <b>Título:</b> {metadata['title']}<br>
        <b>Tipo:</b> {metadata['type']}<br>
        <b>Tamaño:</b> {metadata['size'] / 1024 / 1024:.2f} MB<br>
        <b>Caracteres:</b> {metadata['content_length']:,}<br>
        <br>
        <b>Resumen:</b><br>
        {metadata['summary'][:500]}...
        """
        self.book_info.setHtml(info_text)
        
        # Mostrar conceptos
        self.concepts_list.clear()
        for concept in metadata.get('key_concepts', []):
            self.concepts_list.addItem(f"• {concept}")
            
        # Habilitar botón eliminar
        self.btn_remove.setEnabled(True)
        
    def _remove_book(self):
        """Elimina el libro seleccionado"""
        current_item = self.book_list.currentItem()
        if not current_item:
            return
            
        reply = QMessageBox.question(
            self,
            "Confirmar Eliminación",
            "¿Eliminar este libro de la biblioteca?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            metadata = current_item.data(Qt.ItemDataRole.UserRole)
            
            # Eliminar archivo
            try:
                os.remove(metadata['path'])
            except:
                pass
                
            # Eliminar de la lista
            self.processed_books = [b for b in self.processed_books if b['path'] != metadata['path']]
            self.book_list.takeItem(self.book_list.row(current_item))
            
            # Actualizar
            self._update_stats()
            self._save_library()
            
            # Limpiar detalles
            self.book_info.clear()
            self.concepts_list.clear()
            self.btn_remove.setEnabled(False)
            
    def _update_stats(self):
        """Actualiza las estadísticas"""
        total_books = len(self.processed_books)
        total_size = sum(book['size'] for book in self.processed_books) / 1024 / 1024
        
        self.stats_label.setText(f"📊 Total: {total_books} libros | {total_size:.1f} MB")
        
    def _save_library(self):
        """Guarda el índice de la biblioteca"""
        index_path = os.path.join(self.library_path, 'index.json')
        os.makedirs(os.path.dirname(index_path), exist_ok=True)
        
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(self.processed_books, f, indent=2, ensure_ascii=False)
            
    def _load_library(self):
        """Carga la biblioteca existente"""
        index_path = os.path.join(self.library_path, 'index.json')
        
        if os.path.exists(index_path):
            try:
                with open(index_path, 'r', encoding='utf-8') as f:
                    self.processed_books = json.load(f)
                    
                # Añadir a la lista
                for book in self.processed_books:
                    item = QListWidgetItem(f"📖 {book['title']}")
                    item.setData(Qt.ItemDataRole.UserRole, book)
                    self.book_list.addItem(item)
                    
                self._update_stats()
                
            except Exception as e:
                print(f"Error cargando biblioteca: {e}")
                
    def _notify_ai(self, metadata):
        """Notifica a la IA que hay nuevo conocimiento disponible"""
        # Aquí se integraría con el sistema de IA para que procese el nuevo libro
        self.status_label.setText(
            f"✅ {metadata['title']} añadido. ZEROX está aprendiendo..."
        )